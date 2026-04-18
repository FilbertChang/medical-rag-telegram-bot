import os
import csv
import faiss
import numpy as np
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

FAISS_PATH = "faiss_index"
EMBED_MODEL = "all-MiniLM-L6-v2"
SIMILARITY_THRESHOLD = 0.85

embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)


def parse_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        reader = PdfReader(file_path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext in [".csv"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            return "\n".join(", ".join(row) for row in reader)
    elif ext in [".xls", ".xlsx"]:
        wb = load_workbook(file_path, read_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                rows.append(", ".join(str(c) for c in row if c is not None))
        return "\n".join(rows)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def ingest_file(file_path: str) -> dict:
    # Load existing FAISS index
    db = FAISS.load_local(FAISS_PATH, embeddings, allow_dangerous_deserialization=True)

    # Parse and chunk the new file
    raw_text = parse_file(file_path)
    chunks = splitter.split_documents([Document(page_content=raw_text)])

    added = 0
    updated = 0
    skipped = 0

    for chunk in chunks:
        new_text = chunk.page_content.strip()
        if not new_text:
            continue

        # Search for similar existing chunk
        results = db.similarity_search_with_score(new_text, k=1)

        if not results:
            db.add_documents([Document(page_content=new_text, metadata={"source": os.path.basename(file_path), "specialty": "Uploaded", "description": "User uploaded file"})])
            added += 1
            continue

        top_doc, score = results[0]
        # score in FAISS is L2 distance — lower = more similar
        if score > SIMILARITY_THRESHOLD:
            # No similar chunk found → ADD
            db.add_documents([Document(page_content=new_text, metadata={"source": os.path.basename(file_path), "specialty": "Uploaded", "description": "User uploaded file"})])
            added += 1
        else:
            existing_text = top_doc.page_content.strip()
            if len(new_text) > len(existing_text):
                # New chunk is longer → REPLACE (add new, old stays but gets buried)
                db.add_documents([Document(page_content=new_text, metadata={"source": os.path.basename(file_path), "specialty": "Uploaded", "description": "User uploaded file"})])
                updated += 1
            else:
                # Same or shorter → SKIP
                skipped += 1

    # Save updated index
    db.save_local(FAISS_PATH)

    return {"added": added, "updated": updated, "skipped": skipped}
